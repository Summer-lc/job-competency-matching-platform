from __future__ import annotations

import hashlib
import io
import re
from datetime import date, datetime
from pathlib import Path

from src.skill_ontology import PROFICIENCY_RANK, SKILL_CATALOG, normalize_skill


RESUME_SCHEMA_VERSION = "resume-profile-v2"

SECTION_ALIASES = {
    "basic": ("个人简介", "个人信息", "基本信息", "自我评价", "职业概述", "简介", "summary", "basic information"),
    "skills": ("专业技能", "核心技能", "技能清单", "技能", "skills"),
    "work": ("工作经历", "工作经验", "职业经历", "任职经历", "work experience", "employment"),
    "projects": ("项目经历", "项目经验", "项目", "projects", "project experience"),
    "education": ("教育经历", "教育背景", "学历", "education"),
    "certificates": ("证书资质", "资格证书", "认证证书", "证书", "认证", "certificates", "certifications"),
}

DATE_RANGE_PATTERN = r"""
    (?P<start_year>(?:19|20)\d{2})\s*
    (?:[.\-/]\s*(?P<start_month>\d{1,2})|年\s*(?P<start_month_cn>\d{1,2})\s*月?)
    \s*(?:[-–—~到]|至)\s*
    (?:
        (?P<end_year>(?:19|20)\d{2})\s*
        (?:[.\-/]\s*(?P<end_month>\d{1,2})|年\s*(?P<end_month_cn>\d{1,2})\s*月?)
        |(?P<present>至今|今|现在|present)
    )
"""
DATE_RANGE_RE = re.compile(DATE_RANGE_PATTERN, re.IGNORECASE | re.VERBOSE)

SOURCE_STRENGTH = {
    "work": 0.96,
    "project": 0.93,
    "skill": 0.78,
    "basic": 0.72,
    "certificate": 0.70,
    "education": 0.65,
}

EXCLUDED_SKILL_INTENT_RE = re.compile(
    r"(?:不会|不熟悉|未掌握|未使用|没有使用|计划学习|准备学习|打算学习|将学习|正在学习|学习中|初学|自学|"
    r"not\s+(?:familiar|proficient)|never\s+(?:used?|worked\s+with)|no\s+experience\s+(?:with|in)|"
    r"unfamiliar\s+with|plan(?:ning)?\s+to\s+learn|currently\s+learning|learning)",
    re.IGNORECASE,
)

AFFIRMATIVE_SKILL_INTENT_RE = re.compile(
    r"(?:精通|熟练掌握|熟练使用|掌握|熟悉|使用|开发|负责|建设|主导|"
    r"expert|proficient|familiar|used?|develop(?:ed|ing)?)",
    re.IGNORECASE,
)

PROFICIENCY_MARKERS = (
    ("expert", ("精通", "专家", "expert")),
    ("advanced", ("熟练掌握", "熟练使用", "主导", "架构", "高级", "advanced")),
    ("working", ("掌握", "熟悉", "使用", "开发", "负责", "建设", "工程师", "working")),
)

RESPONSIBILITY_MARKERS = ("负责", "参与", "使用", "开发", "设计", "建设", "维护", "实现", "主导", "优化")
RESULT_MARKERS = ("提升", "降低", "减少", "增长", "节省", "缩短", "达到", "提高", "优化")
UNITLESS_RESULT_MARKERS = ("提升", "降低", "减少", "增长", "节省", "缩短", "达到", "提高")
ROLE_MARKERS = ("工程师", "开发", "架构师", "经理", "负责人", "顾问", "实习生", "技术专家")


def extract_resume_text(raw: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if suffix == ".docx":
        from docx import Document

        document = Document(io.BytesIO(raw))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if text:
                    blocks.append(text)
        return "\n".join(blocks)
    if suffix in {".txt", ".md"}:
        return raw.decode("utf-8-sig", errors="replace").strip()
    raise ValueError("简历仅支持PDF、DOCX、TXT和Markdown格式")


def _clean_line(value: str) -> str:
    return re.sub(
        r"^\s*(?:#{1,6}\s+|[-*•·]\s*|\d{1,2}[、)）]\s*|\d{1,2}\.\s+)",
        "",
        value,
    ).strip()


def _split_sections(text: str) -> tuple[dict[str, list[str]], bool]:
    sections = {name: [] for name in SECTION_ALIASES}
    current = "basic"
    found_heading = False
    aliases = sorted(
        ((alias, section) for section, values in SECTION_ALIASES.items() for alias in values),
        key=lambda item: len(item[0]),
        reverse=True,
    )
    for raw_line in text.splitlines():
        raw_line = raw_line.strip()
        line = _clean_line(raw_line)
        if not line:
            continue
        heading = None
        content = None
        for alias, section in aliases:
            match = re.fullmatch(rf"{re.escape(alias)}\s*(?:(?:[:：])\s*(.*))?", line, re.IGNORECASE)
            if match:
                heading = section
                content = (match.group(1) or "").strip()
                break
        if heading is not None:
            current = heading
            found_heading = True
            if content:
                sections[current].append(content)
            continue
        sections[current].append(raw_line)
    return sections, found_heading


def _reference_date(value: date | datetime | str | None) -> date:
    if value is None:
        return date.today()
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    return date.fromisoformat(str(value)[:10])


def _date_range(value: str, reference_date: date) -> dict | None:
    match = DATE_RANGE_RE.search(value)
    if match is None:
        return None
    start_month = int(match.group("start_month") or match.group("start_month_cn"))
    if match.group("present"):
        end_year, end_month = reference_date.year, reference_date.month
    else:
        end_year = int(match.group("end_year"))
        end_month = int(match.group("end_month") or match.group("end_month_cn"))
    start_year = int(match.group("start_year"))
    if not 1 <= start_month <= 12 or not 1 <= end_month <= 12:
        return None
    start_index = start_year * 12 + start_month - 1
    end_index = end_year * 12 + end_month - 1
    if end_index < start_index:
        return None
    return {
        "match": match,
        "start_date": f"{start_year:04d}-{start_month:02d}",
        "end_date": f"{end_year:04d}-{end_month:02d}",
        "start_index": start_index,
        "end_index": end_index,
        "duration_months": end_index - start_index + 1,
    }


def _skill_occurrence_is_excluded(line: str, start: int) -> bool:
    prefix = re.split(r"[；;。！？!?]|(?:但是|不过|然而|但)", line[:start])[-1]
    excluded = list(EXCLUDED_SKILL_INTENT_RE.finditer(prefix))
    if not excluded:
        return False
    affirmative = list(AFFIRMATIVE_SKILL_INTENT_RE.finditer(prefix))
    return not affirmative or excluded[-1].end() >= affirmative[-1].end()


def _line_skill_mentions(line: str) -> list[dict]:
    mentions: list[dict] = []
    seen = set()
    for catalog_name, (_, catalog_aliases) in SKILL_CATALOG.items():
        aliases = sorted({catalog_name, *catalog_aliases}, key=len, reverse=True)
        for alias in aliases:
            escaped = re.escape(alias)
            if re.fullmatch(r"[A-Za-z0-9_./+ -]+", alias):
                pattern = rf"(?<![A-Za-z0-9_]){escaped}(?![A-Za-z0-9_])"
            else:
                pattern = escaped
            for match in re.finditer(pattern, line, re.IGNORECASE):
                if _skill_occurrence_is_excluded(line, match.start()):
                    continue
                normalized = normalize_skill(alias)
                name = normalized["name"]
                if name and name not in seen:
                    seen.add(name)
                    mentions.append(
                        {
                            "name": name,
                            "category": normalized["category"],
                            "alias": match.group(0),
                        }
                    )
                break
            if catalog_name in seen:
                break
    return mentions


def _skill_names(lines: list[str]) -> list[str]:
    return sorted(
        {mention["name"] for line in lines for mention in _line_skill_mentions(line)},
        key=str.casefold,
    )


def _parse_company_role(header: str) -> tuple[str | None, str | None]:
    value = _clean_line(header).strip(" |｜,，-/")
    if not value:
        return None, None
    if any(value.startswith(marker) for marker in RESPONSIBILITY_MARKERS):
        return None, None
    company_label = re.search(r"(?:公司|单位)\s*[:：]\s*([^|｜,，]+)", value)
    role_label = re.search(r"(?:职位|岗位|职务|角色)\s*[:：]\s*([^|｜,，]+)", value)
    if company_label or role_label:
        return (
            company_label.group(1).strip() if company_label else None,
            role_label.group(1).strip() if role_label else None,
        )
    parts = [part.strip() for part in re.split(r"\s*[|｜]\s*|\s{2,}", value) if part.strip()]
    if (
        len(parts) == 2
        and re.search(r"(?:公司|集团|银行|研究院|研究所|中心|科技)$", parts[0])
        and any(marker in parts[1] for marker in ROLE_MARKERS)
    ):
        return parts[0], parts[1]
    company_match = re.fullmatch(
        r"(?P<company>.+?(?:公司|集团|银行|研究院|研究所|中心|科技))\s+(?P<role>.+)",
        value,
    )
    if company_match:
        role = company_match.group("role").strip()
        if any(marker in role for marker in ROLE_MARKERS):
            return company_match.group("company").strip(), role
    if re.search(r"(?:工程师|架构师|经理|负责人|顾问|实习生|技术专家)$", value):
        return None, value
    return None, None


def _looks_like_work_header(line: str) -> bool:
    company, role = _parse_company_role(line)
    return company is not None or role is not None


def _work_blocks(lines: list[str], reference_date: date) -> list[list[str]]:
    blocks: list[list[str]] = []
    current: list[str] = []
    for line in lines:
        timeline = _date_range(line, reference_date)
        is_header = _looks_like_work_header(line)
        if timeline is not None:
            if current:
                blocks.append(current)
            current = [line]
        elif is_header and len(current) == 1 and _date_range(current[0], reference_date) is not None:
            current.append(line)
        elif is_header:
            if current:
                blocks.append(current)
            current = [line]
        elif current:
            current.append(line)
        else:
            current = [line]
    if current:
        blocks.append(current)
    return blocks


def _parse_work_experiences(lines: list[str], reference_date: date) -> list[dict]:
    experiences = []
    for block in _work_blocks(lines, reference_date):
        timeline = _date_range(block[0], reference_date)
        inline_header = DATE_RANGE_RE.sub("", block[0], count=1).strip() if timeline else ""
        if inline_header and _looks_like_work_header(inline_header):
            header_index, header = 0, inline_header
        elif timeline and len(block) > 1 and _looks_like_work_header(block[1]):
            header_index, header = 1, block[1]
        elif not timeline and _looks_like_work_header(block[0]):
            header_index, header = 0, block[0]
        else:
            header_index, header = None, ""
        company, role = _parse_company_role(header)
        responsibilities = [
            line
            for index, line in enumerate(block)
            if line and index != header_index and not (index == 0 and timeline is not None)
        ]
        experiences.append(
            {
                "company": company,
                "role": role,
                "start_date": timeline["start_date"] if timeline else None,
                "end_date": timeline["end_date"] if timeline else None,
                "duration_months": timeline["duration_months"] if timeline else None,
                "responsibilities": responsibilities,
                "skills": _skill_names(block),
                "evidence_text": "\n".join(block),
            }
        )
    return experiences


def _is_quantified_result(line: str) -> bool:
    quantity = re.search(r"\d+(?:\.\d+)?\s*(?:%|％|倍|万|亿|ms|毫秒|秒|分钟|小时|人日|qps)", line, re.IGNORECASE)
    if quantity is not None and any(marker in line for marker in RESULT_MARKERS):
        return True

    without_noise = re.sub(r"(?<!\d)1[3-9]\d{9}(?!\d)", "", line)
    without_noise = re.sub(
        r"(?:19|20)\d{2}\s*(?:[.\-/年]\s*\d{1,2}(?:\s*月)?)?",
        "",
        without_noise,
    )
    result_pattern = "|".join(map(re.escape, UNITLESS_RESULT_MARKERS))
    return re.search(
        rf"(?:{result_pattern})\s*(?:至|到|为|约)?\s*(?<!\d)\d{{1,9}}(?:\.\d+)?(?!\d)",
        without_noise,
        re.IGNORECASE,
    ) is not None


def _looks_like_project_title(line: str) -> bool:
    line = _clean_line(line)
    return (
        len(line) <= 80
        and not any(marker in line for marker in RESPONSIBILITY_MARKERS)
        and not _is_quantified_result(line)
        and not re.match(r"(?:行业场景|应用场景|项目角色|角色)\s*[:：]", line)
    )


def _project_blocks(lines: list[str], reference_date: date) -> list[list[str]]:
    blocks: list[list[str]] = []
    current: list[str] = []
    has_date = False
    for index, line in enumerate(lines):
        timeline = _date_range(line, reference_date)
        next_has_date = index + 1 < len(lines) and _date_range(lines[index + 1], reference_date) is not None
        explicit_name = re.match(r"(?:项目名称|项目名)\s*[:：]", _clean_line(line)) is not None
        if current and (explicit_name or (has_date and next_has_date and _looks_like_project_title(line))):
            blocks.append(current)
            current = []
            has_date = False
        current.append(line)
        has_date = has_date or timeline is not None
    if current:
        blocks.append(current)
    return blocks


def _parse_project_experiences(lines: list[str], reference_date: date) -> list[dict]:
    projects = []
    for block in _project_blocks(lines, reference_date):
        timeline_lines = [line for line in block if _date_range(line, reference_date) is not None]
        timelines = [_date_range(line, reference_date) for line in timeline_lines]
        timeline = None
        if timelines:
            timeline = {
                "start_date": min(item["start_date"] for item in timelines),
                "end_date": max(item["end_date"] for item in timelines),
            }
        timeline_line = timeline_lines[0] if timeline_lines else None
        header = _clean_line(DATE_RANGE_RE.sub("", timeline_line, count=1)).strip(" |｜,，-/") if timeline_line else ""
        if re.fullmatch(r"(?:项目周期|项目日期|周期|日期|时间)\s*[:：]?", header):
            header = ""
        pre_timeline = block[: block.index(timeline_line)] if timeline_line else block[:1]
        pre_timeline_title = pre_timeline[0] if pre_timeline and _looks_like_project_title(pre_timeline[0]) else ""
        title_source = _clean_line(header or pre_timeline_title)
        parts = [part.strip() for part in re.split(r"\s*[|｜]\s*", title_source) if part.strip()]
        name = parts[0] if parts else None
        role = parts[1] if len(parts) > 1 else None
        if name:
            name = re.sub(r"^(?:项目名称|项目名|项目)\s*[:：]\s*", "", name).strip() or None
        explicit_role = next(
            (
                match.group(1).strip()
                for line in block
                if (match := re.match(r"(?:项目角色|角色)\s*[:：]\s*(.+)", _clean_line(line)))
            ),
            None,
        )
        if explicit_role:
            role = explicit_role
        industry = next(
            (
                match.group(1).strip()
                for line in block
                if (match := re.match(r"(?:行业场景|应用场景)\s*[:：]\s*(.+)", _clean_line(line)))
            ),
            None,
        )
        achievements = [line for line in block if _is_quantified_result(line)]
        metadata = {pre_timeline_title} if pre_timeline_title else set()
        metadata.update(timeline_lines)
        responsibilities = [
            line
            for line in block
            if line not in metadata
            and line not in achievements
            and not re.match(r"(?:行业场景|应用场景|项目角色|角色)\s*[:：]", _clean_line(line))
        ]
        projects.append(
            {
                "name": name,
                "role": role,
                "start_date": timeline["start_date"] if timeline else None,
                "end_date": timeline["end_date"] if timeline else None,
                "skills": _skill_names(block),
                "responsibilities": responsibilities,
                "achievements": achievements,
                "industry_scenario": industry,
                "evidence_text": "\n".join(block),
            }
        )
    return projects


def _proficiency(line: str) -> str:
    lowered = line.casefold()
    for level, markers in PROFICIENCY_MARKERS:
        if any(marker.casefold() in lowered for marker in markers):
            return level
    return "aware"


def _build_skills(
    sections: dict[str, list[str]],
    work_experiences: list[dict],
    project_experiences: list[dict],
) -> list[dict]:
    contexts: list[tuple[str, str, str | None]] = []
    for section, source in (
        ("basic", "basic"),
        ("skills", "skill"),
        ("education", "education"),
        ("certificates", "certificate"),
    ):
        contexts.extend((line, source, None) for line in sections[section])
    for item in work_experiences:
        contexts.extend((line, "work", item["end_date"]) for line in item["evidence_text"].splitlines())
    for item in project_experiences:
        contexts.extend((line, "project", item["end_date"]) for line in item["evidence_text"].splitlines())

    aggregated: dict[str, dict] = {}
    for line, source, used_at in contexts:
        for mention in _line_skill_mentions(line):
            name = mention["name"]
            record = aggregated.setdefault(
                name,
                {
                    "name": name,
                    "category": mention["category"],
                    "aliases": set(),
                    "evidence": [],
                    "proficiency": "aware",
                },
            )
            record["aliases"].add(mention["alias"])
            level = _proficiency(line)
            if PROFICIENCY_RANK[level] > PROFICIENCY_RANK[record["proficiency"]]:
                record["proficiency"] = level
            evidence = {
                "text": line,
                "source": source,
                "strength": SOURCE_STRENGTH[source],
                "used_at": used_at,
            }
            if evidence not in record["evidence"]:
                record["evidence"].append(evidence)

    result = []
    for name, record in aggregated.items():
        evidence = sorted(
            record["evidence"],
            key=lambda item: (-item["strength"], item["source"], item["text"]),
        )
        dated = [item["used_at"] for item in evidence if item["used_at"]]
        sources = list(dict.fromkeys(item["source"] for item in evidence))
        result.append(
            {
                "name": name,
                "category": record["category"],
                "confidence": evidence[0]["strength"],
                "evidence_text": evidence[0]["text"],
                "aliases": sorted(record["aliases"], key=str.casefold),
                "proficiency": record["proficiency"],
                "last_used_at": max(dated) if dated else None,
                "evidence_sources": sources,
                "evidence": evidence,
            }
        )
    return sorted(result, key=lambda item: item["name"].casefold())


def _experience_months(work_experiences: list[dict], sections: dict[str, list[str]]) -> int:
    covered_months = set()
    for item in work_experiences:
        if not item["start_date"] or not item["end_date"]:
            continue
        start_year, start_month = map(int, item["start_date"].split("-"))
        end_year, end_month = map(int, item["end_date"].split("-"))
        start = start_year * 12 + start_month - 1
        end = end_year * 12 + end_month - 1
        covered_months.update(range(start, end + 1))
    if covered_months:
        return len(covered_months)

    fallback_text = "\n".join(sections["basic"] + sections["skills"] + sections["work"])
    values = [
        float(match.group("years"))
        for match in re.finditer(
            r"(?P<years>\d+(?:\.\d+)?)\s*年(?=[^，。；;\n]{0,20}(?:工作|开发|从业|职业|行业|实战)经验)",
            fallback_text,
        )
    ]
    return round(max(values, default=0.0) * 12)


def parse_resume_text(text: str, reference_date: date | datetime | str | None = None) -> dict:
    effective_date = _reference_date(reference_date)
    sections, found_heading = _split_sections(text)
    work_experiences = _parse_work_experiences(sections["work"], effective_date)
    project_experiences = _parse_project_experiences(sections["projects"], effective_date)
    skills = _build_skills(sections, work_experiences, project_experiences)
    experience_months = _experience_months(work_experiences, sections)
    dated_timeline_exists = any(
        item["end_date"] for item in [*work_experiences, *project_experiences]
    )
    cutoff = f"{effective_date.year - 2:04d}-{effective_date.month:02d}"
    if dated_timeline_exists:
        recent_skills = sorted(
            {
                item["name"]
                for item in skills
                if any(
                    evidence["source"] in {"work", "project"}
                    and evidence["used_at"] is not None
                    and evidence["used_at"] >= cutoff
                    for evidence in item["evidence"]
                )
            },
            key=str.casefold,
        )
    else:
        recent_skills = [item["name"] for item in skills]

    education = [degree for degree in ("博士", "硕士", "本科", "大专", "专科") if degree in "\n".join(sections["education"])]
    projects = [item["evidence_text"] for item in project_experiences]
    warnings = [] if found_heading or not text.strip() else ["no_section_headings"]
    return {
        "schema_version": RESUME_SCHEMA_VERSION,
        "parser_mode": "rules",
        "reference_date": effective_date.isoformat(),
        "recency_mode": "dated" if dated_timeline_exists else "undated_fallback",
        "sections": sections,
        "experience_months": experience_months,
        "experience_years": round(experience_months / 12, 2),
        "work_experiences": work_experiences,
        "project_experiences": project_experiences,
        "skills": skills,
        "parse_warnings": warnings,
        "recent_skills": recent_skills,
        "projects": projects,
        "education": education,
        "evidence_count": sum(len(item["evidence"]) for item in skills) + len(projects),
    }


def parse_resume_bytes(
    raw: bytes,
    filename: str,
    reference_date: date | datetime | str | None = None,
) -> dict:
    text = extract_resume_text(raw, filename)
    if not text:
        raise ValueError("未能从简历中提取有效文本")
    parsed = parse_resume_text(text, reference_date=reference_date)
    parsed.update(
        filename=filename,
        content_hash=hashlib.sha256(text.strip().encode("utf-8")).hexdigest(),
        raw_text=text,
    )
    return parsed
