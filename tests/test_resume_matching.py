import io


def _minimal_pdf(text: str) -> bytes:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii"))
        output.extend(obj)
        output.extend(b"\nendobj\n")
    xref = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode("ascii")
    )
    return bytes(output)


def test_docx_resume_parsing_keeps_skill_evidence():
    from docx import Document
    from src.resume_service import parse_resume_bytes

    document = Document()
    document.add_paragraph("张三，4年Java后端开发经验")
    document.add_paragraph("项目经历：使用Java、Spring Boot、MySQL和Docker建设微服务平台。")
    buffer = io.BytesIO()
    document.save(buffer)

    parsed = parse_resume_bytes(buffer.getvalue(), "resume.docx")
    names = {item["name"] for item in parsed["skills"]}
    assert {"Java", "Spring Boot", "MySQL", "Docker"}.issubset(names)
    assert parsed["experience_years"] == 4
    assert parsed["projects"]
    assert all(item["evidence_text"] for item in parsed["skills"])


def test_pdf_resume_parsing_extracts_text():
    from src.resume_service import extract_resume_text

    text = extract_resume_text(_minimal_pdf("Java Python Docker Kubernetes"), "resume.pdf")
    assert "Java" in text
    assert "Kubernetes" in text


def test_matching_returns_dimensions_gaps_and_learning_path():
    from src.matching_service import match_resume_to_job

    resume = {
        "skills": [
            {"name": "Java"},
            {"name": "MySQL"},
            {"name": "Docker"},
        ],
        "recent_skills": ["Java", "MySQL"],
        "experience_years": 4,
        "projects": ["Java微服务平台"],
    }
    job = {
        "name": "Java开发工程师",
        "required_years": 3,
        "required_skills": ["Java", "MySQL", "Kubernetes"],
        "preferred_skills": ["Docker"],
    }
    result = match_resume_to_job(resume, job)
    assert 70 <= result["total_score"] < 100
    assert result["missing_required_skills"] == ["Kubernetes"]
    assert result["learning_path"][0]["skill"] == "Kubernetes"
    assert set(result["dimension_scores"]) == {
        "required_skill_coverage",
        "skill_proficiency",
        "experience_level",
        "project_evidence",
        "skill_recency",
        "preferred_skill_coverage",
        "responsibility_scenario",
    }
